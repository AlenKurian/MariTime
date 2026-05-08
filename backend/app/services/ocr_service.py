import logging
from typing import Tuple
import csv

logger = logging.getLogger(__name__)

TESSERACT_CMD = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def _pytesseract():
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
    return pytesseract


def extract_text_from_pdf(file_path: str) -> Tuple[str, float]:
    # Digital PDF — fast text extraction via PyPDF2
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
        text = "\n".join(text_parts).strip()
        if text:
            return text, 0.95
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")

    # Scanned PDF — render each page to image then OCR with Tesseract
    try:
        import fitz  # PyMuPDF
        from PIL import Image
        import io
        pytesseract = _pytesseract()
        doc = fitz.open(file_path)
        pages_text = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            pages_text.append(pytesseract.image_to_string(img))
        doc.close()
        return "\n\n".join(pages_text).strip(), 0.85
    except Exception as e:
        logger.error(f"Scanned PDF OCR failed: {e}")
        return f"[PDF OCR failed: {e}]", 0.0


def extract_text_from_image(file_path: str) -> Tuple[str, float]:
    try:
        from PIL import Image
        pytesseract = _pytesseract()
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img).strip()
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        confs = [int(c) for c in data["conf"] if str(c).lstrip("-").isdigit() and int(c) >= 0]
        avg_conf = (sum(confs) / len(confs) / 100.0) if confs else 0.85
        return text, avg_conf
    except Exception as e:
        logger.error(f"Image OCR failed: {e}")
        return f"[OCR failed: {e}]", 0.0


def extract_text_from_docx(file_path: str) -> Tuple[str, float, str]:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs), 1.0
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return f"[DOCX extraction failed: {e}]", 0.0


def extract_text_from_odt(file_path: str) -> Tuple[str, float]:
    try:
        from odf.opendocument import load
        from odf.text import P
        from odf import teletype
        doc = load(file_path)
        paragraphs = []
        for elem in doc.text.getElementsByType(P):
            text = teletype.extractText(elem).strip()
            if text:
                paragraphs.append(text)
        return "\n".join(paragraphs), 1.0
    except Exception as e:
        logger.error(f"ODT extraction failed: {e}")
        return f"[ODT extraction failed: {e}]", 0.0


def extract_text_from_csv(file_path: str) -> Tuple[str, float]:
    try:
        rows = []
        with open(file_path, newline="", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(cell.strip() for cell in row))
        return "\n".join(rows), 1.0
    except Exception as e:
        logger.error(f"CSV extraction failed: {e}")
        return f"[CSV extraction failed: {e}]", 0.0


def extract_text(file_path: str, file_type: str) -> Tuple[str, float]:
    extractors = {
        "pdf":  extract_text_from_pdf,
        "jpg":  extract_text_from_image,
        "png":  extract_text_from_image,
        "docx": extract_text_from_docx,
        "odt":  extract_text_from_odt,
        "csv":  extract_text_from_csv,
    }
    extractor = extractors.get(file_type.lower())
    if extractor:
        return extractor(file_path)
    logger.warning(f"No extractor for file type: {file_type}")
    return f"[No extractor available for {file_type}]", 0.0
